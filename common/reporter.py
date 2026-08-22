"""Base report orchestrator — generates HTML, PDF, JSON, CSV from findings."""
from __future__ import annotations

import asyncio
import csv
import io
import json
import logging
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from common.artifact_io import (
    ArtifactBoundaryError,
    atomic_write_bytes,
    ensure_private_directory,
    read_verified_regular_file,
)
from common.confidence_policy import normalise_finding
from common.redaction import redact_text, redact_value, redacted_json_dumps

try:
    from jinja2 import Environment, FileSystemLoader, select_autoescape
    _JINJA2_AVAILABLE = True
except ImportError:
    _JINJA2_AVAILABLE = False

_TEMPLATE_DIR = Path(__file__).parent / "templates"

log = logging.getLogger(__name__)

SEVERITY_ORDER = {"Critical": 0, "High": 1, "Medium": 2, "Low": 3, "Informational": 4}


def _ensure_report_directory(path: Path) -> None:
    """Validate/create a no-follow private directory chain."""
    try:
        ensure_private_directory(path)
    except ArtifactBoundaryError as exc:
        if str(exc) == "artifact directory must be a real directory":
            raise ValueError("report directory must be a real directory") from None
        raise ValueError("report directory is unavailable") from None


def _write_private_bytes(path: Path, content: bytes) -> None:
    """Atomically replace a report through a pinned no-follow parent."""
    atomic_write_bytes(path, content, mode=0o600)


def _write_private_text(path: Path, content: str) -> None:
    """Atomically replace a redacted ordinary artifact with owner-only mode."""
    _write_private_bytes(path, content.encode("utf-8"))


class BaseReporter:
    """Base reporter that serializes findings to multiple output formats."""

    def __init__(
        self,
        findings: list[dict[str, Any]],
        results_dir: Path,
        engagement: str = "default",
        target: str = "",
        tester: str = "anonymous",
        framework: str = "forge",
        formats: list[str] | None = None,
        compliance_collection: Any = None,
        compliance_authority: Any = None,
    ) -> None:
        self.findings     = sorted(
            [normalise_finding(redact_value(f)) for f in findings],
            key=lambda f: (
                -(float(f.get("vpr_score") or f.get("cvss_v31_score") or 0.0)),
                SEVERITY_ORDER.get(f.get("severity", ""), 99),
            ),
        )
        self.results_dir  = Path(results_dir)
        self.engagement   = redact_text(engagement)
        self.target       = redact_text(target)
        self.tester       = redact_text(tester)
        self.framework    = redact_text(framework)
        self.formats      = formats or ["html", "pdf", "json", "csv"]
        self.compliance_collection = compliance_collection
        self.compliance_authority = compliance_authority
        self.generated_at = datetime.now(timezone.utc).isoformat()
        _ensure_report_directory(self.results_dir)

    def generate_all(self) -> dict[str, str]:
        """Generate all configured report formats.

        Returns:
            Dict mapping format name to output file path.
        """
        paths: dict[str, str] = {}
        for fmt in self.formats:
            try:
                method = getattr(self, f"generate_{fmt}", None)
                if method:
                    path = method()
                    if path:
                        paths[fmt] = path
                        log.info("Report generated: %s → %s", fmt, path)
            except Exception:
                log.error("Report generation failed for a configured format")

        try:
            from common.reporting.compliance_engine import ComplianceEngine
            c_reports = ComplianceEngine.evaluate_all(
                self.findings,
                self.compliance_collection,
                self.compliance_authority,
            )
            compliance_path = self.results_dir / "compliance_report.json"
            _write_private_text(
                compliance_path,
                redacted_json_dumps(
                    {k: v.to_dict() for k, v in c_reports.items()}, indent=2
                ),
            )
            paths["compliance"] = str(compliance_path)
            log.info("Compliance report generated: %s", compliance_path)
        except Exception:
            log.warning("Compliance report generation failed")

        return paths

    def _professional_report_config(self, formats: list[str]) -> Any:
        """Build the richer report-engine config lazily to avoid import cycles."""
        from common.reporting.report_engine import ReportConfig
        return ReportConfig(
            engagement=self.engagement,
            target=self.target,
            tester=self.tester,
            output_dir=str(self.results_dir),
            formats=formats,
            include_unverified=False,
            compliance_collection=self.compliance_collection,
            compliance_authority=self.compliance_authority,
        )

    def generate_json(self) -> str:
        """Generate JSON findings export."""
        out = {
            "generated_at": self.generated_at,
            "engagement":   self.engagement,
            "target":       self.target,
            "tester":       self.tester,
            "framework":    self.framework,
            "total":        len(self.findings),
            "summary":      self._severity_summary(),
            "findings":     self.findings,
        }
        path = self.results_dir / "findings.json"
        _write_private_text(path, redacted_json_dumps(out, indent=2, default=str))
        return str(path)

    def generate_docx(self) -> str | None:
        """Generate editable Word report when python-docx is installed."""
        try:
            from docx import Document
        except ImportError:
            log.warning("python-docx not installed — DOCX generation skipped")
            return None

        doc = Document()
        doc.add_heading(f"{self.framework} Penetration Test Report", 0)
        doc.add_paragraph(f"Engagement: {self.engagement}")
        doc.add_paragraph(f"Target: {self.target}")
        doc.add_paragraph(f"Tester: {self.tester}")
        doc.add_paragraph(f"Generated: {self.generated_at}")
        doc.add_heading("Findings", level=1)
        for finding in self.findings:
            doc.add_heading(finding.get("title", "Untitled Finding"), level=2)
            doc.add_paragraph(f"Severity: {finding.get('severity', '')}")
            doc.add_paragraph(f"Confidence: {finding.get('confidence', 'UNVERIFIED')}")
            doc.add_paragraph(f"Verification state: {finding.get('verification_state', 'unknown')}")
            doc.add_paragraph(f"Proof type: {finding.get('proof_type', 'unknown')}")
            doc.add_paragraph(f"Capability maturity: {finding.get('maturity', 'experimental')}")
            doc.add_paragraph(f"Workflow status: {finding.get('status', 'open')}")
            doc.add_paragraph(f"VPR: {finding.get('vpr_score', '')} {finding.get('vpr_priority', '')}")
            doc.add_paragraph(f"Target: {finding.get('url') or finding.get('target', '')}")
            doc.add_paragraph(finding.get("description", ""))
            if finding.get("remediation"):
                doc.add_heading("Remediation", level=3)
                doc.add_paragraph(finding["remediation"])
        path = self.results_dir / "report.docx"
        with io.BytesIO() as artifact:
            doc.save(artifact)
            content = artifact.getvalue()
        _write_private_bytes(path, content)
        return str(path)

    def generate_csv(self) -> str:
        """Generate CSV findings export."""
        path = self.results_dir / "findings.csv"
        fields = [
            "id", "title", "severity", "cvss_v31_score", "target", "port",
            "service", "module", "discovered_at", "confidence", "status",
            "verification_state", "proof_type", "maturity",
        ]
        with io.StringIO(newline="") as artifact:
            writer = csv.DictWriter(artifact, fieldnames=fields, extrasaction="ignore")
            writer.writeheader()
            writer.writerows(self.findings)
            content = artifact.getvalue().encode("utf-8")
        _write_private_bytes(path, content)
        return str(path)

    def generate_html(self) -> str:
        """Generate a self-contained HTML report via Jinja2 template (falls back to inline)."""
        try:
            asyncio.get_running_loop()
        except RuntimeError:
            try:
                from common.reporting.report_engine import ReportEngine
                config = self._professional_report_config(["html"])
                paths = asyncio.run(ReportEngine(self.findings, config).generate())
                if paths.get("html"):
                    return paths["html"]
            except Exception:
                log.debug("Professional report engine unavailable; using fallback HTML")
        else:
            log.debug("Professional report engine unavailable; using fallback HTML")
        path = self.results_dir / "report.html"
        html = self._build_html_jinja2() if _JINJA2_AVAILABLE else self._build_html_inline()
        _write_private_text(path, redact_text(html))
        return str(path)

    def generate_pdf(self) -> str | None:
        """Generate PDF report from HTML via WeasyPrint."""
        try:
            asyncio.get_running_loop()
        except RuntimeError:
            try:
                from common.reporting.report_engine import ReportEngine
                config = self._professional_report_config(["html", "pdf"])
                paths = asyncio.run(ReportEngine(self.findings, config).generate())
                if paths.get("pdf"):
                    return paths["pdf"]
            except Exception:
                log.debug("Professional PDF generation unavailable; using fallback")
        else:
            log.debug("Professional PDF generation unavailable; using fallback")
        try:
            from weasyprint import HTML as WeasyprintHTML
            html_path = self.results_dir / "report.html"
            try:
                html_content = read_verified_regular_file(html_path)
            except ArtifactBoundaryError:
                self.generate_html()
                html_content = read_verified_regular_file(html_path)
            pdf_path = self.results_dir / "report.pdf"
            with io.BytesIO() as artifact:
                WeasyprintHTML(
                    string=html_content.decode("utf-8", errors="replace")
                ).write_pdf(artifact)
                content = artifact.getvalue()
            _write_private_bytes(pdf_path, content)
            return str(pdf_path)
        except ImportError:
            log.warning("WeasyPrint not installed — PDF generation skipped")
            return None
        except Exception:
            log.error("PDF generation failed")
            return None

    def _build_html_jinja2(self) -> str:
        """Render the report via Jinja2 template with auto-escaped output."""
        env = Environment(
            loader=FileSystemLoader(str(_TEMPLATE_DIR)),
            autoescape=select_autoescape(["html", "j2"]),
        )
        template = env.get_template("report.html.j2")

        # Ordinary reports do not embed protected-original screenshots.  Work
        # Package 102 owns authorized evidence-custody and redacted derivatives.
        enriched: list[dict[str, Any]] = []
        for f in self.findings:
            row = dict(f)
            ev = dict(f.get("evidence") or {})
            ev["screenshot_path"] = None
            ev["console_capture_path"] = None
            ev["pcap_path"] = None
            row["evidence"] = ev
            enriched.append(row)

        return template.render(
            framework=self.framework,
            engagement=self.engagement,
            target=self.target,
            tester=self.tester,
            generated_at=self.generated_at,
            findings=enriched,
            summary=self._severity_summary(),
            severity_colors=self._SEVERITY_COLORS,
        )

    def _severity_summary(self) -> dict[str, int]:
        """Return count of findings per severity level."""
        summary: dict[str, int] = {
            "Critical": 0, "High": 0, "Medium": 0, "Low": 0, "Informational": 0
        }
        for f in self.findings:
            sev = f.get("severity", "Informational")
            summary[sev] = summary.get(sev, 0) + 1
        return summary

    _SEVERITY_COLORS: dict[str, str] = {
        "Critical":      "#7030A0",
        "High":          "#FF0000",
        "Medium":        "#FF9900",
        "Low":           "#FFFF00",
        "Informational": "#92D050",
    }

    def _severity_color(self, severity: str) -> str:
        return self._SEVERITY_COLORS.get(severity, "#CCCCCC")

    def _build_html_inline(self) -> str:
        """Build self-contained HTML report string."""
        summary = self._severity_summary()
        rows = ""
        for i, f in enumerate(self.findings, 1):
            sev = f.get("severity", "Informational")
            color = self._severity_color(sev)
            text_color = "#FFFFFF" if sev in ("Critical", "High") else "#000000"
            ev = f.get("evidence", {})
            screenshot_html = ""
            request_html = ""
            if ev.get("request_raw"):
                request_html = (
                    f'<br><strong>Request:</strong><pre style="background:#1e1e1e;color:#d4d4d4;'
                    f'padding:8px;overflow:auto">{self._escape(ev["request_raw"])}</pre>'
                )
            response_html = ""
            if ev.get("response_raw"):
                response_html = (
                    f'<br><strong>Response:</strong><pre style="background:#1e1e1e;color:#d4d4d4;'
                    f'padding:8px;overflow:auto">{self._escape(str(ev["response_raw"])[:2000])}</pre>'
                )
            steps = "".join(
                f"<li>{self._escape(s)}</li>"
                for s in f.get("reproduction_steps", [])
            )
            refs = ", ".join(f.get("references", []))
            mitre = ", ".join(f.get("mitre_attack", []))
            cvss = f.get("cvss_v31_score", "")
            rows += f"""
            <tr>
              <td>{i}</td>
              <td><strong>{self._escape(f.get('title',''))}</strong></td>
              <td style="background:{color};color:{text_color};text-align:center;font-weight:bold">
                {self._escape(sev)}</td>
              <td>{self._escape(str(cvss))}</td>
              <td>{self._escape(f.get('target',''))}</td>
              <td>{self._escape(str(f.get('port','') or ''))}</td>
              <td>{self._escape(f.get('module',''))}</td>
              <td>{self._escape(f.get('verification_state','unknown'))}<br>
                  {self._escape(f.get('proof_type','unknown'))} / {self._escape(f.get('maturity','experimental'))}</td>
            </tr>
            <tr><td colspan="8" style="background:#f9f9f9;padding:12px">
              <strong>Workflow status:</strong> {self._escape(f.get('status','open'))} &nbsp;|&nbsp;
              <strong>Confidence:</strong> {self._escape(f.get('confidence','UNVERIFIED'))}<br>
              <strong>Description:</strong> {self._escape(f.get('description',''))}<br>
              <strong>Reproduction:</strong><ol>{steps}</ol>
              <strong>Remediation:</strong> {self._escape(f.get('remediation',''))}<br>
              <strong>References:</strong> {self._escape(refs)}<br>
              <strong>MITRE ATT&CK:</strong> {self._escape(mitre)}
              {screenshot_html}{request_html}{response_html}
            </td></tr>"""

        summary_bars = ""
        for sev, count in summary.items():
            if count > 0:
                color = self._severity_color(sev)
                tc = "#fff" if sev in ("Critical", "High") else "#000"
                summary_bars += (
                    f'<span style="background:{color};color:{tc};padding:6px 12px;'
                    f'margin:4px;border-radius:4px;font-weight:bold">'
                    f'{sev}: {count}</span>'
                )

        return f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>{self._escape(self.framework)} Report — {self._escape(self.engagement)}</title>
<style>
  body{{font-family:Arial,sans-serif;margin:0;padding:20px;background:#f5f5f5}}
  h1{{color:#333;border-bottom:3px solid #e74c3c;padding-bottom:10px}}
  .meta{{background:#fff;padding:15px;border-radius:6px;margin-bottom:20px;border-left:4px solid #e74c3c}}
  table{{width:100%;border-collapse:collapse;background:#fff;border-radius:6px;overflow:hidden}}
  th{{background:#2c3e50;color:#fff;padding:10px;text-align:left}}
  td{{padding:8px 10px;border-bottom:1px solid #eee;vertical-align:top}}
  tr:hover td{{background:#f0f0f0}}
  pre{{margin:0;white-space:pre-wrap;word-break:break-all;font-size:12px;border-radius:4px}}
  img{{border-radius:4px}}
  .summary{{margin-bottom:15px}}
</style>
</head>
<body>
<h1>{self._escape(self.framework.upper())} — Penetration Test Report</h1>
<div class="meta">
  <strong>Engagement:</strong> {self._escape(self.engagement)} &nbsp;|&nbsp;
  <strong>Target:</strong> {self._escape(self.target)} &nbsp;|&nbsp;
  <strong>Tester:</strong> {self._escape(self.tester)} &nbsp;|&nbsp;
  <strong>Generated:</strong> {self.generated_at} &nbsp;|&nbsp;
  <strong>Total Findings:</strong> {len(self.findings)}
</div>
<div class="summary">{summary_bars}</div>
<table>
  <thead><tr>
    <th>#</th><th>Vulnerability</th><th>Severity</th><th>CVSS</th>
    <th>Target</th><th>Port</th><th>Module</th><th>Verification / Proof</th>
  </tr></thead>
  <tbody>{rows}</tbody>
</table>
<p style="color:#999;font-size:11px;margin-top:20px">
  Generated by {self._escape(self.framework)} — FOR AUTHORIZED USE ONLY
</p>
</body></html>"""

    @staticmethod
    def _escape(text: str) -> str:
        return (text
                .replace("&", "&amp;")
                .replace("<", "&lt;")
                .replace(">", "&gt;")
                .replace('"', "&quot;"))


class TestReporter:
    """Unit tests for reporter module."""

    def _sample_findings(self) -> list[dict]:
        return [{
            "id": "abc123", "title": "SQL Injection", "severity": "High",
            "cvss_v31_score": 8.8, "target": "https://example.com",
            "port": 443, "service": "https", "module": "sqli_scanner",
            "description": "SQLi found", "reproduction_steps": ["Step 1"],
            "remediation": "Use parameterized queries", "references": ["CWE-89"],
            "mitre_attack": ["TA0001/T1190"], "discovered_at": "2024-01-01T00:00:00Z",
            "evidence": {}, "operator_confirmed": False, "tags": [],
            "confidence": "HIGH", "verification_state": "verified",
            "proof_type": "response-confirmed", "maturity": "verified",
        }]

    def test_generate_json(self, tmp_path: Path) -> None:
        r = BaseReporter(self._sample_findings(), tmp_path, formats=["json"])
        paths = r.generate_all()
        assert "json" in paths
        data = json.loads(Path(paths["json"]).read_text())
        assert data["total"] == 1

    def test_generate_csv(self, tmp_path: Path) -> None:
        r = BaseReporter(self._sample_findings(), tmp_path, formats=["csv"])
        paths = r.generate_all()
        assert "csv" in paths
        content = Path(paths["csv"]).read_text()
        assert "SQL Injection" in content

    def test_generate_html(self, tmp_path: Path) -> None:
        r = BaseReporter(self._sample_findings(), tmp_path, formats=["html"])
        paths = r.generate_all()
        assert "html" in paths
        content = Path(paths["html"]).read_text()
        assert "SQL Injection" in content
        assert "<!DOCTYPE html>" in content
